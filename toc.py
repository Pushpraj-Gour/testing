from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches
import test_data
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

bookmark_id_counter = 0  # Global counter for unique bookmarks

def add_bookmark(paragraph, bookmark_name):

    global bookmark_id_counter
    unique_bookmark_name = f"{bookmark_name[:25]}_{bookmark_id_counter}"  # Add unique identifier
    bookmark_start = OxmlElement('w:bookmarkStart')
    bookmark_start.set(qn('w:id'), str(bookmark_id_counter))
    bookmark_start.set(qn('w:name'), unique_bookmark_name)

    bookmark_end = OxmlElement('w:bookmarkEnd')
    bookmark_end.set(qn('w:id'), str(bookmark_id_counter))

    paragraph._p.append(bookmark_start)
    paragraph._p.append(bookmark_end)

    bookmark_id_counter += 1
    return unique_bookmark_name

def format_text(doc, text, style=None, bold=False):

    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    if style:
        paragraph.style = style
    if bold:
        run.bold = True
    return paragraph

def add_horizontal_line(doc):
    paragraph = doc

    p = paragraph._p  
    pPr = p.get_or_add_pPr() 
    pbdr = OxmlElement('w:pBdr')  
    
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')  
    bottom.set(qn('w:sz'), '4')  
    bottom.set(qn('w:space'), '1')  
    bottom.set(qn('w:color'), 'auto')
    
    pbdr.append(bottom)  
    pPr.append(pbdr)  

def add_list_item(doc, list_text, bullet=True):

    para = ""
    for idx,text in enumerate(list_text):
        para+= str(idx+1) + ". " + text + '\n\n' 
    doc.add_paragraph(para)
        


def add_heading(doc, text, level, toc_entries):

    if level == 1:
        doc.add_page_break()

    if len(text) > 100:
        doc.add_paragraph(text)
    else:
        text = text.replace("_", " ").capitalize()
        heading = doc.add_heading(text, level=level)
        bookmark_name = text.replace(' ', '_').lower()
        unique_bookmark_name = add_bookmark(heading, bookmark_name)  # Use unique name

        toc_entries.append((text, level, unique_bookmark_name))  # Store unique name in TOC entries
        if level in {0, 1}:
            add_horizontal_line(heading)

def add_hyperlink_url(paragraph, url, text, color="0000FF", underline=True):

    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')

    rPr = OxmlElement('w:rPr')

    if color:
        c = OxmlElement('w:color')
        c.set(qn('w:val'), color)
        rPr.append(c)

    if underline:
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)

    new_run.append(rPr)
    text_element = OxmlElement('w:t')
    text_element.text = text
    new_run.append(text_element)

    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)


def add_hyperlink_toc(paragraph, bookmark_name, text):
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('w:anchor'), bookmark_name)  # Use unique bookmark name here

    new_run = OxmlElement('w:r')
    text_element = OxmlElement('w:t')
    text_element.text = text
    new_run.append(text_element)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def table_of_contents(doc, toc_entries, target_heading):
    for paragraph in doc.paragraphs:
        if target_heading.lower()==paragraph.text.lower():
         
            toc_paragraph = doc.add_paragraph("Table of Contents\n", style='Heading 1')
            toc_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            paragraph._element.addprevious(toc_paragraph._element)

            headings = []
            for text, level, bookmark_name in toc_entries:
                # if text.lower()=="insights":
                #     continue
                toc_item = doc.add_paragraph()
                if level == 0:

                    toc_item.add_run(f"* ")
                    toc_item.paragraph_format.left_indent = Inches(0.0)

                elif level == 1:

                    toc_item.add_run("• ")
                    toc_item.paragraph_format.left_indent = Inches(0.3)
                elif level == 2:

                    toc_item.add_run("- ")
                    toc_item.paragraph_format.left_indent = Inches(0.6)
                    # continue
                else:
                    continue

                add_hyperlink_toc(toc_item, bookmark_name, text)
                paragraph._element.addprevious(toc_item._element)
                
            
            paragraph._element.addprevious(doc.add_page_break()._element)
            break


def parse_data(doc, data, level,toc_entries):
    if isinstance(data, dict):
        heading = 0 
        for key, value in data.items():

            if isinstance(value, list) or isinstance(value,dict):
                add_heading(doc, key, level,toc_entries)
                parse_data(doc, value,level + 1,toc_entries)
            else:
                if heading ==0:
                    add_heading(doc,value,level,toc_entries)
                    heading+=1
                else:
                    format_text(doc,value)
    
    elif isinstance(data, list):
        # needed_indexing = True
        for item in data:
            if isinstance(item, dict) or isinstance(item, list):
                parse_data(doc,item,level+1,toc_entries)
            else:
                add_list_item(doc,item)

def create_document_from_data(data, file_name='output.docx'):

    doc = Document()

    # toc_paragraph = doc.add_paragraph()
    # insert_table_of_contents(toc_paragraph)

    # Adding Front Page Manually
    title = doc.add_heading()
    add_horizontal_line(title)
    run = title.add_run("Kurlon Report")
    run.font.size = Pt(30)
    company_name = doc.add_paragraph("By Unmistakably Human ")
    add_hyperlink_url(company_name, 'https://www.unmistakablyhuman.ai', "(unmistakablyhuman.ai)")
    add_horizontal_line(company_name)
    doc.add_page_break()


    #TOC thing
    toc_entries = []
    parse_data(doc, data,0,toc_entries)
    table_of_contents(doc,toc_entries,"Key insights")
    doc.save(file_name)
    print(f'Document saved as {file_name}')



import json 

with open('final_output.json', 'r') as file:
    data = json.load(file)

create_document_from_data(data)

