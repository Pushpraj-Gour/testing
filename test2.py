from docx import Document

# Data (similar to your structure)
# data = [
#     {"title": "Kurlon Report"},
#     {
#         "Summary": {
#             "Customer Profiling and Segmentation": '''The user research interviews for Kurlon reveal a brand that successfully caters to a diverse customer base while maintaining a strong focus on sleep quality and comfort. Several key themes emerge from the findings:
            
# 1. **Broad Market Appeal**: Kurlon has positioned itself as a versatile brand, appealing to a wide range of demographics from middle to upper-middle-class urban families. This diversity is reflected in their product pricing, which spans from budget-friendly to premium options, allowing the brand to capture various income segments within the Indian market.

# 2. **Customer Lifestyle and Aspirations**: Kurlon's customers exhibit varied lifestyles, from career-focused individuals to those prioritizing home improvement and travel. This diversity in customer profiles presents both opportunities and challenges for targeted marketing and product development.

# 3. **Technology and Digital Engagement**: There's a clear trend towards digital platforms for entertainment and information among Kurlon customers. This insight offers potential avenues for enhanced digital marketing strategies and customer engagement.
# '''
#         }
#     },
#     {
#         "Brand Perceptions and Positioning": '''Kurlon emerges as a well-established and trusted mattress brand in the Indian market, with a strong reputation built on a foundation of reliability, affordability, and quality.'''
#     }
# ]


data = [
    {"title": "Kurlon Report"},
    {
        "Summary": {
            "Customer Profiling and Segmentation": """The user research interviews for Kurlon reveal a brand that successfully caters to a diverse customer base while maintaining a strong focus on sleep quality and comfort. Several key themes emerge from the findings:

1. **Broad Market Appeal**: Kurlon has positioned itself as a versatile brand, appealing to a wide range of demographics from middle to upper-middle-class urban families. This diversity is reflected in their product pricing, which spans from budget-friendly to premium options, allowing the brand to capture various income segments within the Indian market.

2. **Customer Lifestyle and Aspirations**: Kurlon's customers exhibit varied lifestyles, from career-focused individuals to those prioritizing home improvement and travel. This diversity in customer profiles presents both opportunities and challenges for targeted marketing and product development.

3. **Technology and Digital Engagement**: There's a clear trend towards digital platforms for entertainment and information among Kurlon customers. This insight offers potential avenues for enhanced digital marketing strategies and customer engagement.

4. **Sleep Quality Focus**: Customers consistently prioritize comfort, support, and durability in their mattresses, with medium-firm options being most popular.

5. **Health and Wellness Awareness**: There's a growing consciousness among customers about mattress hygiene, temperature regulation, and orthopedic support. This awareness aligns well with Kurlon's focus on health benefits and technological advancements in their products.

6. **Demographic Influences**: Factors such as age, income, family structure, occupation, and geographic location significantly impact customer preferences and purchasing decisions.
""",
        }
    },
    {
        "Brand Perceptions and Positioning": """Kurlon emerges as a well-established and trusted mattress brand in the Indian market, with a strong reputation built on a foundation of reliability, affordability, and quality. The brand's success stems from its ability to cater to a wide range of customer segments while maintaining a consistent image of value and trustworthiness. Key strengths of the Kurlon brand include broad appeal across diverse demographics, emotional connections based on trust and comfort, superior value for money, wide availability, and a diverse product range catering to various needs and preferences.
""",
    },
    {
        "Customer Profiling and Segmentation": {
            "Customer Demographics Analysis": """Kurlon caters to a broad spectrum of customers, with its core demographic being middle to upper-middle-class urban families. The brand appeals to both young couples starting their families and established households with school-age children. Customers typically have stable, white-collar jobs or run small to medium businesses, with education levels ranging from graduate to postgraduate degrees. Kurlon successfully targets different income segments, offering products that cater to budget-conscious consumers as well as those seeking premium options.""",
            "Key Takeaways": """Kurlon's customer base is diverse, including young couples, established families, IT professionals, teachers, business owners, and corporate professionals. The pricing range from 12,000 to 50,000+ INR reflects the brand's versatility in meeting the needs of different demographics.""",
        }
    },
    {
        "Customer Lifestyle and Aspirations": {
            "Key Takeaways": """The lifestyle patterns, aspirations, and future goals of Kurlon customers are diverse and multifaceted. They range from career-focused individuals balancing work and family life to those prioritizing home improvement and travel experiences. Housing situations vary from joint family setups to modern apartments, reflecting the diverse socio-economic backgrounds of Kurlon's customer base. Entertainment preferences show a strong inclination towards digital platforms, with social media and OTT services playing a significant role in daily life.""",
            "Synthesized Insights": """Kurlon customers exhibit diverse lifestyle patterns, often influenced by work, family, and health considerations. Aspirations include career growth, financial stability, quality education for children, home ownership, and travel experiences. Housing situations and entertainment preferences reflect socio-economic diversity, with a mix of joint families, nuclear households, modern apartments, and traditional homes.""",
        }
    },
    {
        "Purchase Decision-Making and Channel Preferences": """The user research for Kurlon reveals a complex interplay between online and offline channels throughout the customer journey, with distinct preferences and behaviors across different demographic segments. Key themes include the importance of comfort, the role of price and brand reputation, the omnichannel research behavior, and the critical role of in-store experiences.
""",
    },
    {
        "Product Experience, Expectations, and Development": """Kurlon customers exhibit diverse preferences for comfort, support, and features, influenced by factors like age, health, and lifestyle. The brand faces challenges with heat retention, adjustment periods, and limited waterproofing in some products. There are opportunities to expand into sleep-related products like pillows and protective bedding, and to innovate in advanced materials and health features.
""",
    },
    {
        "Market Positioning and Strategic Opportunities": """Kurlon's broad product range and competitive pricing resonate well with middle and upper-middle-class consumers. The brand should continue to emphasize its core values while adapting its messaging to address the unique preferences of different customer segments. Opportunities lie in innovation, sustainability, targeted marketing, and customer experience enhancement.
""",
    }
]

# Initialize Document
doc = Document()

# Add the title from the first dictionary
doc.add_heading(data[0]["title"], level=0)

# Iterate through the remaining sections in the 'data' list
for section in data[1:]:
    for heading, content in section.items():
        # Add section heading
        doc.add_heading(heading, level=1)
        
        # If the content is another dictionary (like the 'Summary' section), iterate through it
        if isinstance(content, dict):
            for subheading, text in content.items():
                doc.add_heading(subheading, level=2)
                doc.add_paragraph(text)
        else:
            # Otherwise, it's just a string (like 'Brand Perceptions and Positioning'), add the content
            doc.add_paragraph(content)

# Save the document
doc.save('kurlon_report.docx')
