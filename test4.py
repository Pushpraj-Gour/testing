from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def format_text(doc, text, style=None, bold=True):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    if style:
        paragraph.style = style
    if bold:
        run.bold = True
    return paragraph

def add_list_item(doc, text, bullet=False, reset_numbering=True):
    """Add list items - bullet or numbered. Reset numbering if required."""
    # Choose style based on bullet or numbered list
    style = 'List Bullet' if bullet else 'List Number'

    # Reset numbering by creating a new paragraph list if requested
    paragraph = doc.add_paragraph(text, style=style)
    if reset_numbering and not bullet:
        # Restart numbering using custom numbering style
        set_numbering_restart(paragraph)

def set_numbering_restart(paragraph):
    """Resets numbering for a paragraph to start from 1."""
    p = paragraph._element
    pPr = p.get_or_add_pPr()
    numPr = OxmlElement('w:numPr')
    ilvl = OxmlElement('w:ilvl')
    ilvl.set(qn('w:val'), "0")
    numId = OxmlElement('w:numId')
    numId.set(qn('w:val'), "1")
    numPr.append(ilvl)
    numPr.append(numId)
    pPr.append(numPr)

def add_heading(doc, text, level):
    """Add heading to the document with a max level of 3."""
    level = min(level, 3)
    doc.add_heading(text, level=level)

def parse_data(doc, data, level=-1):
    """Recursive function to parse data structure and create a formatted document."""
    if isinstance(data, dict):
        for key, value in data.items():
            add_heading(doc, key, level=level)

            # When encountering a new heading or sub-heading, reset numbering
            if isinstance(value, dict) or isinstance(value, list):
                parse_data(doc, value, level=level + 1)
            else:
                format_text(doc, str(value))

    elif isinstance(data, list):
        # Reset numbering for each list under a different heading
        reset_numbering = True
        for item in data:
            if isinstance(item, dict) or isinstance(item, list):
                parse_data(doc, item, level=level + 1)
            else:
                # Use the reset_numbering flag to restart numbering for new lists
                add_list_item(doc, str(item), bullet=False, reset_numbering=reset_numbering)
                reset_numbering = False  # Only reset numbering for the first item in a new list

def create_document_from_data(data, file_name='output.docx'):
    doc = Document()

    # Parse the data structure
    parse_data(doc, data)

    # Save the document
    doc.save(file_name)
    print(f'Document saved as {file_name}')


import test_data

create_document_from_data(test_data.data)
