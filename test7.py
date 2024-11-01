data = "\nKurlon's customer base is diverse and multifaceted, encompassing middle to upper-middle-class urban families across various life stages and income levels. This diversity is reflected in the brand's wide price range and product offerings, allowing it to cater to both budget-conscious consumers and those seeking premium options.\n\n### Key Themes from Research\n\n1. **Product Performance and Sleep Quality**  \n   Customers prioritize a balance of comfort, support, and durability in mattresses, with medium-firm options being the most popular. Those who switched to Kurlon generally report improved sleep quality, highlighting the brand's effectiveness in addressing common sleep issues such as physical discomfort and stress-related disturbances.\n\n2. **Holistic Sleep Environment**  \n   There is a growing awareness among customers about the importance of the overall sleep environment, including factors like temperature regulation, hygiene, and complementary accessories. This presents an opportunity for Kurlon to expand its product line and educate customers on comprehensive sleep solutions.\n\n3. **Demographic Influences**  \n   Age, life stage, income, family structure, occupation, and geographic location significantly impact customer preferences and purchasing decisions. This underscores the need for targeted marketing strategies and product development to address specific segment needs.\n\n4. **Technology and Lifestyle Integration**  \n   Kurlon's customers exhibit diverse technology usage and social media habits, suggesting potential for tailored digital marketing and engagement strategies across different platforms.\n\n5. **Brand Perception**  \n   Kurlon is viewed as a versatile brand capable of meeting varied needs, with associations of technological advancement and health benefits, particularly among urban professionals.\n\n### Strategic Recommendations for Kurlon\n\nTo capitalize on these insights, Kurlon should consider:\n\n- **Developing targeted marketing campaigns** that address specific demographic needs.\n- **Expanding its product range** to cater to evolving customer requirements.\n- **Strengthening its brand positioning** around sleep quality and overall well-being.\n- **Enhancing its digital presence** and engagement strategies.\n- **Educating customers** on the importance of a holistic sleep environment.\n\n### Future Research Directions\n\nFurther research could explore:\n\n- The interaction between demographic and psychographic factors.\n- Quantitative analysis of digital platform usage.\n- Potential untapped market segments for growth opportunities."

from docx import Document
from docx.shared import Pt

doc = Document()

# list_data = data.split("\n")
# for item in list_data:
#     paragraph = doc.add_paragraph(item)
#     paragraph_format = paragraph.paragraph_format
#     paragraph_format.space_after = Pt(0)  # Set space after to 0
#     paragraph_format.space_before = Pt(0)  # Set space before to 0

sentences = [
    "- **Enhancing its digital presence** and engagement strategies.",
    "Increasing market share through innovative marketing techniques.",
    "- **Boosting customer satisfaction** by improving service quality.",
    "hello **my** name is **pushpraj**",
    "Hello **how** **are** you and **I** am fine **what about** you"
]

# for sent in sentences:
#     list_sent = sent.split("**")
#     print(list_sent,"\n")


def add_sentence_with_bold(paragraph, sentence):
    # Split the sentence by '**', which marks the bold text
    parts = sentence.split("**")
    
    for i, part in enumerate(parts):
        # Even indices are non-bold parts, odd indices are bold parts
        if i % 2 == 0:
            paragraph.add_run(part)  # Regular text
        else:
            bold_run = paragraph.add_run(part)
            bold_run.bold = True  # Bold text
for sentence in sentences:
    # Add a new paragraph for each sentence
    paragraph = doc.add_paragraph()
    
    # Handle bold and regular text
    add_sentence_with_bold(paragraph, sentence)

# Save the document
doc.save("formatted_sentences.docx")

